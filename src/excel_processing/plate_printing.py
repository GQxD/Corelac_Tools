import os
import re
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def ajouter_bordures_cellule(cell):
    """Ajoute des bordures Ã  une cellule de tableau Word"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    
    tcPr.append(tcBorders)

def extraire_numero_plaque(nom_fichier):
    """Extrait le numÃ©ro de la plaque depuis le nom du fichier"""
    match = re.search(r'Plaque_(\d{3})', nom_fichier)
    if match:
        return int(match.group(1))
    return None

def determiner_temperature(numero_plaque):
    """DÃ©termine la tempÃ©rature en fonction du numÃ©ro de plaque"""
    if numero_plaque <= 100:
        return "5Â°C"
    elif numero_plaque <= 200:
        return "9Â°C"
    else:
        return "N/A"

def ajouter_plaque_dans_cellule(cell, nom_plaque, ws, numero_plaque):
    """Ajoute le titre et le tableau d'une plaque dans une cellule EN CONSERVANT LE FORMATAGE"""
    # DÃ©terminer la tempÃ©rature
    temperature = determiner_temperature(numero_plaque)
    
    # Ajouter le titre de la plaque avec la tempÃ©rature
    titre_para = cell.add_paragraph()
    titre_run = titre_para.add_run(f"{nom_plaque} - {temperature}")
    titre_run.bold = True
    titre_run.font.size = Pt(11)
    titre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Espace minimal
    cell.add_paragraph()
    
    # Obtenir les dimensions du tableau Excel
    max_row = ws.max_row
    max_col = ws.max_column
    
    # CrÃ©er le tableau dans la cellule
    table = cell.add_table(rows=max_row, cols=max_col)
    table.style = 'Light Grid Accent 1'
    
    # Remplir le tableau EN CONSERVANT LE FORMATAGE EXCEL
    for i, row in enumerate(ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col)):
        for j, excel_cell in enumerate(row):
            word_cell = table.rows[i].cells[j]
            
            # Obtenir la valeur de la cellule
            valeur = excel_cell.value
            if valeur is None:
                valeur = ""
            else:
                valeur = str(valeur)
            
            word_cell.text = valeur
            
            # IMPORTANT: RÃ©cupÃ©rer le formatage Excel (gras, italique)
            excel_font = excel_cell.font
            est_gras_excel = excel_font.bold if excel_font and excel_font.bold else False
            est_italique_excel = excel_font.italic if excel_font and excel_font.italic else False
            
            # Formater la cellule Word
            for paragraph in word_cell.paragraphs:
                for run in paragraph.runs:
                    if i == 0:  # En-tÃªte
                        run.font.bold = True
                        run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Appliquer la taille de base
                        run.font.size = Pt(8)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # CONSERVER LE FORMATAGE EXCEL
                        if est_gras_excel:
                            run.font.bold = True
                        if est_italique_excel:
                            run.font.italic = True
            
            # Ajouter des bordures
            ajouter_bordures_cellule(word_cell)

def ajouter_legende(doc):
    """Ajoute une page de lÃ©gende au dÃ©but du document"""
    # Titre
    titre = doc.add_heading('LÃ©gende des plaques', level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Section TempÃ©ratures
    temp_heading = doc.add_heading('TempÃ©ratures des plaques', level=2)
    temp_para1 = doc.add_paragraph()
    temp_para1.add_run('Plaque_001 Ã  Plaque_100 : ').bold = True
    temp_para1.add_run('5Â°C')
    
    temp_para2 = doc.add_paragraph()
    temp_para2.add_run('Plaque_101 Ã  Plaque_200 : ').bold = True
    temp_para2.add_run('9Â°C')
    
    doc.add_paragraph()
    
    # Section Formatage
    format_heading = doc.add_heading('Formatage des Ã©chantillons', level=2)
    
    # Gras
    gras_para = doc.add_paragraph()
    gras_run = gras_para.add_run('Texte en gras')
    gras_run.bold = True
    gras_para.add_run(' : Femelles modifiÃ©es (B_F11-F15 remplacÃ©es)')
    
    # Italique
    italic_para = doc.add_paragraph()
    italic_run = italic_para.add_run('Texte en italique')
    italic_run.italic = True
    italic_para.add_run(' : Interversion B_M1 â†” B_M7')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Note informative
    note_para = doc.add_paragraph()
    note_para.add_run('Note : ').bold = True
    note_para.add_run('Ces modifications ont Ã©tÃ© appliquÃ©es pour optimiser la distribution des Ã©chantillons.')
    
    # Saut de page aprÃ¨s la lÃ©gende
    doc.add_page_break()

def regrouper_plaques_word(dossier_source, dossier_destination=None, plaques_par_page=6):
    """
    CrÃ©e UN document Word avec plusieurs plaques par page (format 2x3).
    CONSERVE LE FORMATAGE (gras, italique) des fichiers Excel.
    
    Args:
        dossier_source: Chemin du dossier contenant les fichiers Excel
        dossier_destination: Chemin du dossier pour le fichier Word
        plaques_par_page: Nombre de plaques par page (dÃ©faut: 6)
    """
    # VÃ©rifier que le dossier source existe
    if not os.path.exists(dossier_source):
        print(f"âŒ ERREUR: Le dossier '{dossier_source}' n'existe pas!")
        return
    
    # Lister tous les fichiers du dossier
    print(f"ðŸ“ Contenu du dossier:")
    tous_fichiers = os.listdir(dossier_source)
    
    # Filtrer uniquement les fichiers Plaque_XXX
    fichiers_excel = []
    for f in tous_fichiers:
        if re.match(r'Plaque_\d{3}\.(xlsx|xls)$', f, re.IGNORECASE):
            fichiers_excel.append(f)
    
    fichiers_excel.sort()
    
    print(f"   Total de fichiers: {len(tous_fichiers)}")
    print(f"   Fichiers Plaque_XXX trouvÃ©s: {len(fichiers_excel)}")
    
    if len(fichiers_excel) == 0:
        print("\nâš ï¸  Aucun fichier Plaque_XXX.xlsx trouvÃ©!")
        return
    
    print(f"\n   Fichiers Ã  inclure:")
    for f in fichiers_excel[:5]:
        print(f"   - {f}")
    if len(fichiers_excel) > 5:
        print(f"   ... et {len(fichiers_excel) - 5} autres")
    
    nb_pages = (len(fichiers_excel) + plaques_par_page - 1) // plaques_par_page
    print(f"\n   Configuration: {plaques_par_page} plaques par page")
    print(f"   ðŸ“„ Nombre de pages: {nb_pages + 1} (incluant la lÃ©gende)")
    print(f"   âœ¨ Conservation du formatage Excel (gras, italique)\n")
    
    # CrÃ©er le dossier de destination
    if dossier_destination is None:
        dossier_destination = os.path.join(dossier_source, "documents_word")
    
    if not os.path.exists(dossier_destination):
        os.makedirs(dossier_destination)
    
    # CrÃ©er le document Word
    doc = Document()
    
    # Ajouter la page de lÃ©gende en premier
    ajouter_legende(doc)
    
    # Configuration en mode paysage et marges rÃ©duites
    sections = doc.sections
    for section in sections:
        section.orientation = 1  # Paysage
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)
    
    plaques_ajoutees = 0
    fichiers_erreur = 0
    
    # Traiter les plaques par groupes de 6 (2 lignes x 3 colonnes)
    for i in range(0, len(fichiers_excel), plaques_par_page):
        batch = fichiers_excel[i:i+plaques_par_page]
        
        print(f"CrÃ©ation page {i//plaques_par_page + 2}/{nb_pages + 1}...")
        
        # CrÃ©er un tableau 2x3 pour contenir 6 plaques
        main_table = doc.add_table(rows=2, cols=3)
        main_table.autofit = False
        main_table.allow_autofit = False
        
        # DÃ©finir la largeur des colonnes (plus larges sans les notes)
        for col in main_table.columns:
            for cell in col.cells:
                cell.width = Inches(3.5)
        
        # DÃ©finir la hauteur des lignes (plus hautes)
        main_table.rows[0].height = Inches(4.0)
        main_table.rows[1].height = Inches(4.0)
        
        # Position des plaques dans le tableau 2x3
        positions = [(0, 0), (0, 1), (0, 2), (1, 0), (1, 1), (1, 2)]
        
        for idx, fichier in enumerate(batch):
            chemin_fichier = os.path.join(dossier_source, fichier)
            nom_plaque = os.path.splitext(fichier)[0]
            numero_plaque = extraire_numero_plaque(fichier)
            
            try:
                # Charger le fichier Excel avec data_only=False pour conserver le formatage
                wb = load_workbook(chemin_fichier, data_only=True)
                ws = wb.worksheets[0]
                
                # Obtenir la cellule du tableau principal
                row_idx, col_idx = positions[idx]
                cell = main_table.rows[row_idx].cells[col_idx]
                
                # Ajouter la plaque dans cette cellule (avec formatage conservÃ©)
                ajouter_plaque_dans_cellule(cell, nom_plaque, ws, numero_plaque)
                
                plaques_ajoutees += 1
                print(f"  âœ“ {nom_plaque} ajoutÃ© avec formatage ({plaques_ajoutees}/{len(fichiers_excel)})")
                
            except Exception as e:
                print(f"  âœ— Erreur avec {fichier}: {str(e)}")
                fichiers_erreur += 1
        
        # Ajouter un saut de page sauf pour la derniÃ¨re page
        if i + plaques_par_page < len(fichiers_excel):
            doc.add_page_break()
    
    # Sauvegarder le document Word
    if plaques_ajoutees > 0:
        nom_word = f"Plaques_condensÃ©es_{plaques_ajoutees}_plaques_{nb_pages + 1}_pages_FORMATÃ‰.docx"
        chemin_word = os.path.join(dossier_destination, nom_word)
        doc.save(chemin_word)
        
        print("\n" + "="*60)
        print(f"Export terminÃ© !")
        print(f"âœ“ Plaques incluses: {plaques_ajoutees}")
        print(f"âœ— Fichiers en erreur: {fichiers_erreur}")
        print(f"ðŸ“„ Pages crÃ©Ã©es: {nb_pages + 1} (incluant lÃ©gende)")
        print(f"ðŸ’¾ Ã‰conomie: {len(fichiers_excel) - nb_pages} pages de plaques")
        print(f"âœ¨ Formatage Excel conservÃ© (gras, italique)")
        print(f"ðŸ“ Document sauvegardÃ©: {chemin_word}")
        print("="*60)
    else:
        print("\nâŒ Aucune plaque n'a pu Ãªtre ajoutÃ©e au document.")


if __name__ == "__main__":
    # Dossier source attendu: contient des fichiers Excel de plaques (Plaque_XXX.xlsx)
    dossier_source = r"A_REMPLACER_PAR_CHEMIN_DOSSIER"
    
    # Optionnel: dossier de sortie Word (sinon un sous-dossier local sera créé automatiquement)
    # dossier_destination = r"A_REMPLACER_PAR_CHEMIN_DOSSIER"
    
    print("DÃ©but de la crÃ©ation du document Word condensÃ© (v2 avec formatage)...")
    print(f"Dossier source: {dossier_source}\n")
    
    regrouper_plaques_word(dossier_source)
    
    print("\nAppuyez sur EntrÃ©e pour fermer...")
    input()
